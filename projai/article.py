"""
    This module provides a simple API for structuring
    article data in map-like data structure and parse
    it as json or docx.

"""

from typing import List

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Content:
    
    def __init__(self, ) -> None:
        self.elements: List[dict] = []
    
    def __getitem__(self, index):
        return self.elements[index]

    def __repr__(self):
        return repr(self.elements)

    def add_subsection(self, title: str) -> None:
        self.elements.append({"subsection": title})

    def add_subsubsection(self, title: str) -> None:
        self.elements.append({"subsubsection": title})

    def add_paragraph(self, text: str) -> None:
        self.elements.append({"paragraph": text})

    def add_paragraphs(self, texts: List[str]) -> None:
        self.elements.extend([{"paragraph": text} for text in texts])
    
    def add_unordered_list(self, items: List[str]):
        self.elements.append({"unordered_list": items})

    def add_ordered_list(self, items: List[str]):
        self.elements.append({"ordered_list": items})

    def add_table(self, table: dict):
        header = table['header']
        rows = table['rows']
        if not all(len(header) == len(row) for row in rows):
            raise ValueError('Rows must have same length as header')
        self.elements.append(
            {
                "table": {"thead": header, "tbody": rows},
            }
        )


class Section:
    def __init__(self, title: str = None) -> None:
        self.title = title
        self.content: Content = Content()

    @property
    def elements(self,) -> list:
        return self.content.elements

    def json(self,) -> dict:
        return {
            "title": self.title,
            "content": self.content.elements
        }


class DocumentProcessor:

    def __init__(self, title: str = None) -> None:
        self.title = title
        self.sections: List[Section] = []

    def add_section(self, title: str = None) -> Section:
        self.sections.append(Section(title))
        return self.sections[-1]
    
    def add_content(self,
                    *arg,
                    content_type: str = 'paragraph',
                    **kwargs):
        if len(self.sections) == 0:
            self.add_section()
        method = f"add_{content_type}"
        content = self.sections[-1].content
        try:
            getattr(content, method)(*arg, **kwargs)
        except AttributeError as err:
            methods = [
                method.rsplit('add_', maxsplit=1)[-1]
                for method in dir(content) if 'add_' in method
            ]
            message = (
                f"Invalid content type '{content_type}'. "
                f"Valids are: {methods}"
            )
            raise ValueError(message) from err

    def json(self,) -> dict:
        return {
            "title": self.title,
            "sections": [section.json() for section in self.sections]
        }

    def load(self, source: dict):
        self.title = source["title"]
        for section in source["sections"]:
            try:
                self.add_section(section["title"])
            except KeyError as err:
                raise err
            key = 'content'
            if 'sections' in section:
                key = 'sections'
            for content in section[key]:
                for element in content:
                    if element == "content":
                        for sub_element in content[element]:
                            for item in sub_element:
                                self.add_content(
                                    sub_element[item],
                                    content_type=item
                                )
                        continue
                    self.add_content(content[element], content_type=element)
    
    def docx(self,) -> Document:
        """
        Generates a Word document based on the content of the
        RawDocument object.
        
        Args:
            output_file (str): Path to save the generated Word document.
        """
        document = Document()
        document.add_heading(self.title, 0)

        for section in self.sections:
            document.add_heading(section.title, level=1)
            for content in section.content.elements:
                for content_type, data in content.items():
                    if content_type == 'subsection':
                        document.add_heading(data, level=2)
                    elif content_type == 'subsubsection':
                        document.add_heading(data, level=3)
                    elif content_type == 'paragraph':
                        paragraph = document.add_paragraph(data)
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragraph_format.first_line_indent = Inches(0.5)
                    elif content_type == 'unordered_list':
                        for item in data:
                            document.add_paragraph(item, style='List Bullet')
                    elif content_type == 'ordered_list':
                        for item in data:
                            document.add_paragraph(item, style='List Number')
                    elif content_type == 'table':
                        table = document.add_table(
                            rows=1, cols=len(data['thead'])
                        )
                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(data['thead']):
                            paragraph = hdr_cells[i].add_paragraph()
                            paragraph.add_run(header).bold = True
                        for row_data in data['tbody']:
                            row_cells = table.add_row().cells
                            for i, cell_data in enumerate(row_data):
                                row_cells[i].text = cell_data
                        document.add_paragraph("")

        return document
